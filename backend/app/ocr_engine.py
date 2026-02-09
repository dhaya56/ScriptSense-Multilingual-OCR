import os
import cv2
import numpy as np
import tempfile
import uuid
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont # Import TTFont directly here
import traceback # Import traceback for detailed error logging

from app.translator import translate_text, normalize_lang
from app.utils import (
    detect_language,
    get_doc_stats,
    highlight_low_confidence_words,
    LANGUAGE_MAP,
    LANG_FONT_MAP,
    correct_spelling,
    grammar_correction,
    detect_handwritten_or_printed, # This is imported and will be used in preprocess_image
    extract_images_from_docx
)
from app.ocr_router import extract_text_with_best_model

# ------------------- Preprocessing -------------------

def preprocess_image(image_path): # Removed is_handwritten flag from here
    """
    Preprocesses an image by converting to grayscale, blurring, applying adaptive or OTSU thresholding
    based on whether the image is detected as handwritten or printed, and deskewing it.
    
    Args:
        image_path (str): The path to the input image.

    Returns:
        str: The path to the saved preprocessed image.
    """
    image = cv2.imread(image_path)
    if image is None:
        raise ValueError(f"Could not read image at {image_path}")

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)

    # Detect if the image is handwritten or printed to inform thresholding
    is_handwritten_local = detect_handwritten_or_printed(image_path) 

    # --- Adaptive Thresholding for Handwritten vs. Printed ---
    if is_handwritten_local:
        # For handwritten text, adaptive thresholding often works better
        # It handles varying lighting conditions across the image
        thresh = cv2.adaptiveThreshold(blurred, 255, 
                                       cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                       cv2.THRESH_BINARY_INV, 11, 2)
        print(f"[INFO] Applying adaptive thresholding for handwritten image: {os.path.basename(image_path)}")
    else:
        # For printed text, OTSU is usually very effective
        _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        print(f"[INFO] Applying OTSU thresholding for printed image: {os.path.basename(image_path)}")

    # --- Deskewing Logic ---
    deskewed = thresh # Initialize in case coords is empty
    try:
        coords = np.column_stack(np.where(thresh < 128)) # Use a lower threshold for coordinates, looking for darker pixels
        if coords.size == 0:
            print(f"[WARN] No strong foreground pixels found for deskewing in {os.path.basename(image_path)}. Skipping deskew operation.")
        else:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            
            (h, w) = thresh.shape
            M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
            deskewed = cv2.warpAffine(thresh, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
            print(f"[INFO] Deskewed image {os.path.basename(image_path)} by {angle:.2f} degrees.")
    except Exception as e:
        print(f"[ERROR] Deskewing failed for {os.path.basename(image_path)}: {e}")
        traceback.print_exc()
        deskewed = thresh # Fallback to un-deskewed thresholded image

    # --- Save Preprocessed Image ---
    base, ext = os.path.splitext(image_path)
    # Ensure a unique name for the preprocessed file
    processed_path = f"{base}_preprocessed{ext}"
    cv2.imwrite(processed_path, deskewed)
    print(f"[INFO] Saved preprocessed image to: {processed_path}")
    return processed_path

# ------------------- OCR Handlers -------------------

def handle_image(file_path, enhance=True, source_lang='en'):
    if enhance:
        file_path = preprocess_image(file_path)
    text, word_conf, low_conf = extract_text_with_best_model(file_path, source_lang=source_lang)
    return text, word_conf, low_conf

def handle_pdf(file_path, enhance=True, source_lang='en'):
    images = convert_from_path(file_path)
    all_text, word_conf, low_conf = "", [], []

    for idx, img in enumerate(images):
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
        img.save(tmp_path, "PNG")

        try:
            path_to_process = preprocess_image(tmp_path) if enhance else tmp_path # preprocess_image now handles its own handwritten detection
            text, wc, lowc = extract_text_with_best_model(path_to_process, source_lang=source_lang)
            if text.strip():
                all_text += f"\n[Page {idx + 1}]\n{text}"
            word_conf.extend(wc)
            low_conf.extend(lowc)
        except Exception as e:
            print(f"[ERROR] Failed to process PDF page {idx + 1}: {e}")
            traceback.print_exc()
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path) # Clean up original temp image
            if enhance and 'path_to_process' in locals() and path_to_process != tmp_path and os.path.exists(path_to_process):
                os.unlink(path_to_process) # Clean up preprocessed temp image

    return all_text.strip(), word_conf, low_conf

def handle_docx(file_path, enhance=True, source_lang='en'):
    doc = Document(file_path)
    text_parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    extracted_text = "\n".join(text_parts)

    images = extract_images_from_docx(file_path)
    # For text already present in DOCX, assign a high confidence as it's not OCR'd
    word_conf = [(w, 0.99) for w in extracted_text.split()] 
    low_conf = []

    for idx, img_path in enumerate(images):
        try:
            path_to_process = preprocess_image(img_path) if enhance else img_path # preprocess_image now handles its own handwritten detection
            img_text, wc, lowc = extract_text_with_best_model(path_to_process, source_lang=source_lang)
            if img_text.strip():
                extracted_text += f"\n[Image {idx + 1}]\n{img_text}"
            word_conf.extend(wc)
            low_conf.extend(lowc)
        except Exception as e:
            print(f"[ERROR] Failed to process DOCX image {idx + 1}: {e}")
            traceback.print_exc()
        finally:
            # extract_images_from_docx uses a TemporaryDirectory, which handles cleanup.
            # No explicit unlink needed here as the temp_dir context manager will remove them.
            pass

    return extracted_text.strip(), word_conf, low_conf

# ------------------- File Routing -------------------

def extract_text_dynamic(file_path, enhance=True, source_lang='en'):
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ['.jpg', '.jpeg', '.png']:
        return handle_image(file_path, enhance, source_lang)
    elif ext == '.pdf':
        return handle_pdf(file_path, enhance, source_lang)
    elif ext == '.docx':
        return handle_docx(file_path, enhance, source_lang)
    return "Unsupported file format", [], []

# ------------------- Output Writers -------------------

def save_to_pdf(text, path, lang_code='en'):
    font_name = "Helvetica" 
    font_size = 11
    
    lang_code = normalize_lang(lang_code or "en")
    font_path = LANG_FONT_MAP.get(lang_code, LANG_FONT_MAP.get("en")) 
    
    print(f"[DEBUG] Font path for {lang_code}: {font_path}")
    print(f"[DEBUG] Font file exists: {os.path.exists(font_path) if font_path else 'N/A'}")

    if font_path and os.path.exists(font_path):
        try:
            custom_font_name = f"Noto_{lang_code}" 
            if custom_font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(custom_font_name, font_path))
            font_name = custom_font_name
        except Exception as e:
            print(f"[ERROR] Failed to register font for {lang_code}: {e}")
            traceback.print_exc()
            font_name = "Helvetica" 
    else:
        print(f"[WARN] Font not found for language: {lang_code} at {font_path}, using default 'Helvetica'.")
        font_name = "Helvetica" 

    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4
    x_margin, y_margin = inch * 0.75, inch * 0.75
    y = h - y_margin
    
    if font_name in pdfmetrics.getRegisteredFontNames() or font_name in ["Helvetica", "Times-Roman", "Courier"]:
        c.setFont(font_name, font_size)
    else:
        print(f"[WARN] Requested font '{font_name}' is not registered. Using 'Helvetica'.")
        c.setFont("Helvetica", font_size)

    for para in text.split("\n"):
        words = para.split(" ")
        line = ""
        for word in words:
            if c.stringWidth(line + word + " ", font_name, font_size) < w - 2 * x_margin:
                line += word + " "
            else:
                c.drawString(x_margin, y, line.strip())
                y -= font_size + 3 
                line = word + " "
                
                if y < y_margin:
                    c.showPage()
                    y = h - y_margin
                    if font_name in pdfmetrics.getRegisteredFontNames() or font_name in ["Helvetica", "Times-Roman", "Courier"]:
                        c.setFont(font_name, font_size)
                    else:
                        c.setFont("Helvetica", font_size)
        
        if line:
            c.drawString(x_margin, y, line.strip())
            y -= font_size + 3
        
        y -= 5 # Add extra spacing between paragraphs

        if y < y_margin:
            c.showPage()
            y = h - y_margin
            if font_name in pdfmetrics.getRegisteredFontNames() or font_name in ["Helvetica", "Times-Roman", "Courier"]:
                c.setFont(font_name, font_size)
            else:
                c.setFont("Helvetica", font_size)
                
    c.save()

def save_to_docx(text, path):
    doc = Document()
    font = doc.styles['Normal'].font
    font.name = 'Calibri'
    font.size = Pt(11)

    for para in text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.save(path)

# ------------------- Main Pipeline -------------------

def process_document(file_path, source_lang=None, target_lang=None, enhance=False):
    # Determine the language to use for OCR based on source_lang or detection
    # Ensure source_lang is normalized for consistency
    initial_source_lang = normalize_lang(source_lang) if source_lang else None

    # Perform OCR
    extracted_text, word_conf, low_conf_words = extract_text_dynamic(file_path, enhance, source_lang=initial_source_lang)
    
    if not extracted_text.strip():
        print("[ERROR] No text could be extracted from the document.")
        return {"error": "No text extracted"}

    # Detect language if not explicitly provided
    detected_lang_code, confidence = detect_language(extracted_text)
    
    # Use provided source_lang or detected language for further processing
    lang_to_use = initial_source_lang or detected_lang_code 
    print(f"[INFO] Using language for processing: {lang_to_use} (Detected: {detected_lang_code}, Provided: {initial_source_lang})")

    # Spell correction (only for 'en' or 'hi' as per your logic)
    extracted_text_corrected = correct_spelling(extracted_text, lang_code=lang_to_use)

    # Normalize target language code
    target_lang_code = normalize_lang(target_lang or 'en') # Default to 'en' if target_lang is None

    # Translate text
    translated = translate_text(extracted_text_corrected, lang_to_use, target_lang_code)
    
    # Apply spell correction to translated text (only for 'en' or 'hi' as per your logic)
    translated_corrected = correct_spelling(translated, lang_code=target_lang_code)
    
    # Apply grammar correction to translated text (only for 'en' as per your logic)
    translated_final = grammar_correction(translated_corrected, lang_code=target_lang_code)

    # Get document statistics
    stats = get_doc_stats(extracted_text_corrected, file_path, chars_per_line=80)

    file_id = str(uuid.uuid4())
    results_dir = "static/results"
    os.makedirs(results_dir, exist_ok=True) # Ensure results directory exists

    # Save extracted text
    extracted_pdf_path = os.path.join(results_dir, f"extracted_{file_id}.pdf")
    extracted_docx_path = os.path.join(results_dir, f"extracted_{file_id}.docx")
    save_to_pdf(extracted_text_corrected, extracted_pdf_path, lang_code=lang_to_use)
    save_to_docx(extracted_text_corrected, extracted_docx_path)

    # Save translated text
    translated_pdf_path = os.path.join(results_dir, f"translated_{file_id}.pdf")
    translated_docx_path = os.path.join(results_dir, f"translated_{file_id}.docx")
    save_to_pdf(translated_final, translated_pdf_path, lang_code=target_lang_code)
    save_to_docx(translated_final, translated_docx_path)

    avg_conf = sum([conf for _, conf in word_conf]) / len(word_conf) if word_conf else 1.0

    print("[DEBUG] Word Confidence Scores (sample):", word_conf[:3])  # sample of (word, conf)
    print("[DEBUG] Flattened Confidence Scores (sample):", [conf for _, conf in word_conf[:3]])

    return {
        "extracted_text": highlight_low_confidence_words(extracted_text_corrected, low_conf_words),
        "translated_text": translated_final,
        "detected_language": detected_lang_code,
        "confidence": confidence, # Confidence of language detection
        "stats": stats,
        "low_conf_count": len(low_conf_words),
        "word_confidence_scores": [conf for _, conf in word_conf],
        "confidence_metrics": {
            "document_quality": min(100, round(avg_conf * 100 + (10 if enhance else 0))), # Add bonus for enhancement
            "handwriting_clarity": round(avg_conf * 100), # This metric might be misleading if text is printed
            "text_recognition": round(avg_conf * 100) # Use OCR average confidence for recognition
        },
        "download_extracted_pdf": f"/{extracted_pdf_path}", 
        "download_extracted_docx": f"/{extracted_docx_path}",
        "download_translated_pdf": f"/{translated_pdf_path}",
        "download_translated_docx": f"/{translated_docx_path}"
    }